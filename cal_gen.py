import calendar
import os
import csv_gen as csvgen
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from settings import MONTH_CODES, MONTH_COLORS, WEEK_COLORS, WEEK_CODES, ROOMS, WEEK_STARTS_ON, CAL_FONT_SIZE, UNDERLINE_START_TIME, UNDERLINE_END_TIME, UNDERLINE_BOOL
from datetime import date, datetime
from settings import NUMBER_OF_ROOMS, CORRECTION
import settings as SETTINGS


def create_document(cal_type,input_date,output_folder):
    if cal_type == "Week":
        return create_week(input_date,output_folder)
        
    if cal_type == "Month":
        return create_month(input_date,output_folder)
    

def create_week(input_date,output_folder):

    TYPE_OF_CAL = "week"

    ''' create weekly calendar using the chosen monday as an input in MM/DD/YYYY format '''

    # get user input and turn it into usable data and strings 
    in_month = int(input_date[1])
    in_day = int(input_date[2])
    in_year = int(input_date[0])

    # user number input -> string english output
    num_days = calendar.monthrange(in_year,in_month)[1] 

    # generate csv file from input
    csv_weekBegin = f"{in_month}/{in_day}/{in_year}"
    
    # create range of dates and header/file name
    if determine_day(in_day+6,num_days) < in_day:
        # if the week ends in a new month or year, name file and header appropriately
        if in_month + 1 == 13:
            # if the last week containing dates from december goes into january, change date to be january of next year 
            csv_weekEnd = f"{1}/{determine_day(in_day+6,num_days)}/{in_year+1}"
            month_year_day = f'Week of {calendar.month_name[in_month]} {in_day} {in_year} to {calendar.month_name[1]} {determine_day(in_day+6,num_days)} {in_year+1}'
        else:
            csv_weekEnd = f"{in_month+1}/{determine_day(in_day+6,num_days)}/{in_year}"
            month_year_day = f'Week of {calendar.month_name[in_month]} {in_day} to {calendar.month_name[in_month+1]} {determine_day(in_day+6,num_days)} {in_year}'
        new_month = True
    
    else:
        # if the week ends in the same month, name file and header appropriately also 
        csv_weekEnd = f"{in_month}/{determine_day(in_day+6,num_days)}/{in_year}"
        month_year_day = f'Week of {calendar.month_name[in_month]} {in_day} to {in_day+6} {in_year}'
        new_month = False

    # generate CSV file and get filename
    print("Getting events from .ical file... ", end="")
    csv_filename = csvgen.calgen(csv_weekBegin,csv_weekEnd,TYPE_OF_CAL)
    print("Done!")

    # set up document properties
    print("Setting up document... ", end="")
    document, table = setup_document(month_year_day,TYPE_OF_CAL)
    print("Done!")

    # draw table based on calendar for filling 
    print("Adding Monday through Sunday to header... ", end="")
    add_dates_to_header(table,in_day,num_days)
    print("Done!")
    
    # get events in array form
    print("Getting and converting events into array... ", end="")
    week_events = get_events(csv_filename,in_day,num_days,TYPE_OF_CAL)
    print("Done!")

    # add events from array
    print("Adding events to table cells:\n")
    add_events_to_week(table,week_events)
    print("Done adding events!")

    # create file name and save file
    document_name = month_year_day + '.docx'
    document_name = os.path.join(output_folder, document_name)

    document.save(document_name)
    
    return document_name, csv_filename

    
def add_dates_to_header(table,first_day,num_days):

    ''' creates calendar layout as a word table. ''' 
    
    header_cells = table.rows[0].cells

    days_list = []
    if WEEK_STARTS_ON == "Monday":
        days_names = ['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
    elif WEEK_STARTS_ON == "Sunday":
        days_names = ['Sunday','Monday','Tuesday','Wednesday','Thursday','Friday','Saturday']

    for i in range(7):
        n = first_day+i
        number = determine_day(n,num_days)
        number_day = f'{days_names[i]} {number}'
        days_list.append(number_day)

    for day in range(7):
        header_cells[day+1].text = days_list[day]
        p = header_cells[day+1].paragraphs[0]
        header_cells[day+1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_format = p.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_grey_bg = parse_xml(r'<w:shd {} w:fill="f2f2f2"/>'.format(nsdecls('w')))
        header_cells[day+1]._tc.get_or_add_tcPr().append(cell_grey_bg)
        
def add_events_to_week(table,week_events):
    # give cells slightly grey background
    cell_grey_bg = parse_xml(r'<w:shd {} w:fill="f2f2f2"/>'.format(nsdecls('w')))
    table.cell(0,0)._tc.get_or_add_tcPr().append(cell_grey_bg)
    for row in range(1,9):
        # determine color for each row
        colors = WEEK_COLORS[row]
        color = RGBColor(colors[0],colors[1],colors[2])
        
        print(f"Adding {ROOMS[row-1]} events...")
        
        for column in range(8):
            # establish cell to use and its alignment
            cell = table.cell(row,column)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # reset this variable on each cell, as each cell has a pre-existing blank paragraph
            first_event = True
            
            # add room name to first column
            if column == 0:
                cell_grey_bg = parse_xml(r'<w:shd {} w:fill="f2f2f2"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(cell_grey_bg)
                if row != 0:
                    p = cell.paragraphs[0]
                    p_properties = p.add_run(ROOMS[row-1])
                    p_properties.font.color.rgb = color
                    p_properties.bold =  True
                    p_properties.font.size = Pt(CAL_FONT_SIZE)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
            
            # add events to the following columns 
            else:
                day_events = week_events[column-1]
                sort_list_of_lists(day_events,2)
                for event in range(len(day_events)):
                    
                    if day_events[event][1] == WEEK_CODES[row]:
                        
                        # establish strings
                        event_name, event_start_time, event_end_time = event_strings(day_events[event],WEEK_CODES,"week")

                        # determine if event is an evening event, if option enabled
                        if int(day_events[event][3]) >= UNDERLINE_START_TIME and int(day_events[event][3]) <= UNDERLINE_END_TIME and UNDERLINE_BOOL == True:
                            is_underlined = True
                        else:
                            is_underlined = False                    

                        fontsize = CAL_FONT_SIZE
            
                        # add event times and determine the style
                        if first_event == True:
                            # if this is the first line of the cell, make sure to use pre-existing first paragraph so there is no blank first line in each cell
                            time_string = f"{event_start_time} - {event_end_time}"
                            time = cell.paragraphs[0]
                            time_properties = time.add_run(time_string)
                            time_properties.font.color.rgb = color
                            time_properties.bold =  False
                            time_properties.underline = is_underlined
                            time_properties.font.size = Pt(fontsize)
                            time.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # create new paragraph if this is not the first line
                            time_string = f"{event_start_time} - {event_end_time}"
                            time = cell.add_paragraph()
                            time_properties = time.add_run(time_string)
                            time_properties.font.color.rgb = color
                            time_properties.bold =  False
                            time_properties.underline = is_underlined
                            time_properties.font.size = Pt(fontsize)
                            time.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # add event name and determine the style
                        event = cell.add_paragraph()
                        event_properties = event.add_run(event_name)
                        event_properties.font.color.rgb = color
                        event_properties.bold =  True
                        event_properties.underline = is_underlined
                        event_properties.font.size = Pt(fontsize)
                        event.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        first_event = False

def create_month(input_date,output_folder):

    TYPE_OF_CAL = "month"

    input_date.pop(2)
    # get user input and turn it into usable data and strings 
    in_month = int(input_date[1])
    in_year = int(input_date[0])
    start_of_month = calendar.weekday(in_year,in_month,1)

    # convert datetime weekday numbers to that of the corresponding column 
    cal_to_table = {6 : 0, 0 : 1, 1 : 2, 2 : 3, 3 : 4, 4 : 5, 5 : 6,}

    # gives the day of the week the month starts on
    start_of_month = cal_to_table[start_of_month]

    # user number input -> string english output
    month_and_year = f'{calendar.month_name[in_month]} {in_year}'
    
    # get number of days in month
    num_days = calendar.monthrange(in_year,in_month)[1] + 1

    # generate csv file from input
    csv_monthBegin = f"{in_month}/1/{in_year}"
    csv_monthEnd = f"{in_month}/{num_days-1}/{in_year}"
    calendar.monthrange(in_year,in_month)[1]
    print("\nGetting events from .ical... ", end="")
    csv_filename = csvgen.calgen(csv_monthBegin,csv_monthEnd,TYPE_OF_CAL)
    print("Done!")

    # set up document properties
    print("Setting up document... ", end="")
    document, table = setup_document(month_and_year,TYPE_OF_CAL)
    print("Done!")

    # draw table based on calendar for filling 
    print("Adding Sunday through Saturday to header and calendar dates to cells... ", end="")
    add_dates_to_table(start_of_month,table,num_days)
    print("Done!")

    # get events in array form
    print("Turning events into array... ", end="")    
    month_events = get_events(csv_filename,1,num_days,TYPE_OF_CAL)
    print("Done!")

    # add events from array
    print("Adding events from array into table... ", end="")    
    add_events_to_month(start_of_month,table,num_days,month_events)
    print("Done!")
    
    # create file name and save file
    document_name = month_and_year + '.docx'
    document_name = os.path.join(output_folder, document_name)
    document.save(document_name)

    return document_name, csv_filename
 
def add_dates_to_table(start_of_month,table,num_days):

    ''' creates calendar layout as a word table. ''' 
    
    def add_date_to_cell(table,row,column,counter_days):
        cell = table.cell(row,column)
        cell_grey_bg = parse_xml(r'<w:shd {} w:fill="f7f7f7"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(cell_grey_bg)
        p = cell.paragraphs[0]
        p_properties = p.add_run(str(counter_days))
        p_properties.bold =  True
        p_properties.font.size = Pt(CAL_FONT_SIZE)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        counter_days += 1
        return counter_days

    header_cells = table.rows[0].cells
    days_list = ['Sunday','Monday','Tuesday','Wednesday','Thursday','Friday','Saturday']
    for day in range(7):
        header_cells[day].text = days_list[day]

    # month starts on a sunday
    if start_of_month == 0:
        counter_days = 1
        for row in range(1,14,2):
            for column in range(7):
                if counter_days == num_days:
                    break
                counter_days = add_date_to_cell(table,row,column,counter_days)
    
    # month starts on a saturday
    elif start_of_month == 6:
        counter_days = 1
        for row in range(1,16,2):
            if row == 1:
                counter_days = add_date_to_cell(table,row,6,counter_days)
            else:
                for column in range(7):
                    if counter_days == num_days:
                        break
                    counter_days = add_date_to_cell(table,row,column,counter_days)


    # month starts on a week day
    elif start_of_month in [1,2,3,4,5]:
        counter_days = 1
        for row in range(1,14,2):
            if counter_days == num_days:
                break
            if row == 1:
                for column in range(start_of_month,7):
                    counter_days = add_date_to_cell(table,row,column,counter_days)
            else:
                for column in range(7):
                    if counter_days == num_days:
                        break
                    counter_days = add_date_to_cell(table,row,column,counter_days)
        
def add_events_to_month(start_of_month,table,num_days,month_events):

    def add_events_to_cell(cell,month_events,counter_days):

        # cell parameters
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # sort day's events by time 
        sort_list_of_lists(month_events[counter_days-1],2)   

        number_of_events = 0

        for event in month_events[counter_days-1]:
            # to be used, eventually, to figure out how to set the font size automatically based on how many events there are in a given cell
            number_of_events += 1
            
            # create paragraph 
            p = cell.paragraphs[0]
            p_format = p.paragraph_format
            p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # get strings for assembling into readable event
            
            event_name, event_start_time, event_end_time, event_room= event_strings(event,MONTH_CODES,"month")
            
            # determine if event is an evening event
            if int(event[3]) > 1800:
                is_underlined = True
            else:
                is_underlined = False                    

            # add event to cell
            add_info_text = p.add_run(f'{event_room} {event_start_time}-{event_end_time} ')
            if event == month_events[counter_days-1][-1]:
                add_event_name = p.add_run(f'{event_name}')
            else:
                add_event_name = p.add_run(f'{event_name}\n')                    
                    
            # set font styles
            red = MONTH_COLORS[event_room][0]
            green = MONTH_COLORS[event_room][1]
            blue = MONTH_COLORS[event_room][2]
            font_size = CAL_FONT_SIZE
            color = RGBColor(red,green,blue)
            add_event_name.bold = True
            add_event_name.font.size = Pt(font_size)
            add_info_text.font.size = Pt(font_size)
            add_info_text.underline = is_underlined
            add_event_name.underline = is_underlined
            add_event_name.font.color.rgb = color
            add_info_text.font.color.rgb = color

    counter_days = 1

    # if month starts on sunday
    if start_of_month == 0:
        for row in range(2,15,2):
            for column in range(7):
                if counter_days == num_days:
                    break
                cell = table.cell(row,column)
                add_events_to_cell(cell,month_events,counter_days)
                counter_days += 1
    
    # if month starts on saturday
    if start_of_month == 6:
        for row in range(2,17,2):
            if row == 2:
                cell = table.cell(row,6)
                add_events_to_cell(cell,month_events,counter_days)
                counter_days += 1
            else:
                for column in range(7):
                    if counter_days == num_days:
                        break
                    cell = table.cell(row,column)
                    add_events_to_cell(cell,month_events,counter_days)
                    counter_days += 1


    # every week day
    elif start_of_month in [1,2,3,4,5]:
        for row in range(2,17,2):
            if row == 2:
                for column in range(start_of_month,7):
                    cell = table.cell(row,column)
                    add_events_to_cell(cell,month_events,counter_days)
                    counter_days += 1
            else:
                for column in range(7):
                    if counter_days == num_days:
                        break
                    cell = table.cell(row,column)
                    add_events_to_cell(cell,month_events,counter_days)
                    counter_days += 1

def sort_list_of_lists(list_of_list,sublist_position):
    ''' sorts a list of integers that is within a list itself. takes a list with lists in it as the first input, and the array position of the sublist to sort from '''
    
    # this didn't really need to be a function... as it is already a function... but i did this for readability reasons.
    list_of_list.sort(key = lambda x:x[sublist_position])
    return list_of_list

def determine_day(current_day,days_in_month):

    ''' function to determine if a day is in a new month. compares current_day input to days_in_month to see if current_day is higher. returns proper calendar number'''

    if current_day > days_in_month:
        result = abs(days_in_month - current_day)
        return result
    else:
        return current_day
    
def get_events(csv_filename,first_day,num_days,month_or_week):    
    ''' add calendar events to array from source .csv file. arguments are the csv file name, the first day of whatever period is chosen, the number of days in the starting month, and whether or not the calendar is a monthly or weekly one.''' 
    
    period_events = []

    if month_or_week == "week":
        last_day = first_day + 7
    elif month_or_week == "month":
        last_day = num_days

    for day in range(first_day,last_day):
        day_events = []
        with open(csv_filename,'r') as csv_input_file:
            for event in csv_input_file:
                event = event.split("|")
                event[5] = int(event[5])
                
                if month_or_week == "week":
                    # if a week ends in a different month, cheat around the problem with a minorly hacky solution!
                    if event[5] < first_day:
                        event[5] = num_days + event[5]

                # add the event to its respective day array
                if event[5] == day:
                    day_events.append(event)
            
            # add each day's array to the week's array
            period_events.append(day_events)
    return period_events

def event_strings(event,room_codes,month_or_week):

    if month_or_week == "month":
        event_room = room_codes[event[1]]

    event_name = event[0]
    event_start_time = datetime.strptime(event[2], "%H%M")
    event_start_time = event_start_time.strftime("%I:%M%p")
    event_start_time = changeAMPM(event_start_time)
    event_end_time = datetime.strptime(event[3], "%H%M")
    event_end_time = event_end_time.strftime("%I:%M%p")
    event_end_time = changeAMPM(event_end_time)

    if month_or_week == "month":
        return event_name,event_start_time,event_end_time,event_room
    elif month_or_week == "week":
        return event_name,event_start_time,event_end_time

def changeAMPM(var_name):
    ''' input is a string. changes any "AM" to "a", "PM" to "p", and removes 0 if it is the first character ''' 

    if "PM" in var_name:
        var_name = var_name.replace("PM","p")
    if "AM" in var_name:
        var_name = var_name.replace("AM","a")
    if var_name[0] == '0':
        var_name = var_name.replace("0",'',1)
    if ":00" in var_name[-4:]:
        var_name = var_name.replace(":00",'',1)

    return var_name

def setup_document(header_title,month_or_week):
    if month_or_week == "week":
        # TODO: make this number more dynamic to prevent index error
        num_rows = NUMBER_OF_ROOMS - 1
        num_cols = 8
    elif month_or_week == "month":
        num_rows = 13
        num_cols = 7       
    
    document = Document()
    document.add_paragraph()
    p = document.paragraphs[0]
    p_format = p.paragraph_format
    p2 = document.add_paragraph()
    p2_format = p2.paragraph_format
    header_print = p.add_run(f'{header_title}')
    header_print.font.size = Pt(SETTINGS.HEADER_FONT_SIZE)
    today = date.today()
    update_print = p2.add_run(f'Last update: {today.strftime("%B %d, %Y")}')
    p2_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    font = document.styles['Normal'].font
    font.name = SETTINGS.FONT_CHOICE
    font.size = Pt(SETTINGS.CAL_FONT_SIZE)
    
    section = document.sections[0]
    section.page_height = Inches(SETTINGS.PAPER_SIZE_HEIGHT)
    section.page_width = Inches(SETTINGS.PAPER_SIZE_WIDTH)
    section.left_margin = Inches(SETTINGS.MARGIN_SIZE)
    section.right_margin = Inches(SETTINGS.MARGIN_SIZE)
    section.top_margin = Inches(SETTINGS.MARGIN_SIZE)
    section.bottom_margin = Inches(SETTINGS.MARGIN_SIZE)
    
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = SETTINGS.TABLE_STYLE
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    return document, table